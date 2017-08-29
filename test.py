# -*- coding: utf-8 -*-
import docx
import table
import Read
import table_style
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
from reportlab.pdfbase.ttfonts import TTFont
pdfmetrics.registerFont(TTFont('msyh', 'msyh.ttf'))
pdfmetrics.registerFont(TTFont('msyhbd', 'msyhbd.ttc'))
from reportlab.lib import colors
from reportlab.platypus import Table,TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate

READ_TYPE=2

styles = getSampleStyleSheet()
styleH = styles['Heading1']
stylesheet=getSampleStyleSheet()
normalStyle = styles['Normal']
style = stylesheet['Normal']
normalStyle.spaceAfter = 5  #段后距
normalStyle.spaceBefore = 10  #段前距
normalStyle.leading = 20     #行间距
styleH.spaceAfter = 15
styleH.spaceBefore = 15
styleH.leading = 15
styleH.fontName="msyhbd"


def rpt():
    story = []
    #标题：段落的用法详见reportlab-userguide.pdf中chapter 6 Paragraph
    if READ_TYPE == 1:           #读取数据库里的文章
        rpt_title = '<para autoLeading="off" fontSize=10.5 align=center>'+Read.news_title+'<br/></para>'
        story.append(Paragraph(rpt_title, styleH))
        table_list=[]
        temp=[]
        temp.append(Read.news_author)
        temp.append(Read.news_time)
        table_list.append(temp)
        first_table = Table(table_list,colWidths=[225,225])
        first_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (1, 0), 'msyhbd'),
        ('ALIGN', (1, 0), (1, 0), 'RIGHT'),  # 右对齐
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('FONTSIZE', (1, 2), (-1, -1), 10.5),  # 字体大小
        ('TEXTCOLOR', (0, 0), (1, 0), colors.standardfl),
        ]))
        story.append(first_table)
        news_content_list=Read.news_content.splitlines()
        for i in range(len(news_content_list)):
            text = ''
            text += '<para autoLeading="off" fontSize=10.5 ><font face="msyh" >'
            text+=news_content_list[i]
            text +='</font></para>'
            story.append(Paragraph(text, normalStyle))
    if READ_TYPE==2:                  #读取word文档里的文章和表格
        doc = docx.Document('test1.docx')
        length = len(doc.paragraphs)
        rpt_title = '<para autoLeading="off" fontSize=10.5 align=center>' + doc.paragraphs[0].text + '<br/></para>'
        story.append(Paragraph(rpt_title, styleH))
        rpt_title = '<para autoLeading="off" fontSize=10.5 align=right color=rgb(84,141,212)>' + doc.paragraphs[1].text + '<br/></para>'
        story.append(Paragraph(rpt_title, styleH))
        for i in range(2,length):
            text = ''
            if doc.paragraphs[i].style.name=='Heading 3' or doc.paragraphs[i].style.name=='Heading 2'or doc.paragraphs[i].style.name=='Heading 1':
                text+='<para autoLeading="off" fontSize=10.5 ><font face="msyhbd" >'
            else :
                text += '<para autoLeading="off" fontSize=10.5 ><font face="msyh" >'
            text+=doc.paragraphs[i].text
            text +='</font></para>'
            story.append(Paragraph(text, normalStyle))
        #表格数据：用法详见reportlab-userguide.pdf中chapter 7 Table
        table.parse_docx('test1.docx')
        component_data = table.component_data
        component_table = Table(component_data)
        component_table.setStyle(table_style.table_style)
        story.append(component_table)

    doc = SimpleDocTemplate('D:/HX/reportlab/bug1.pdf')
    doc.build(story)

if __name__ == '__main__':
    rpt()



